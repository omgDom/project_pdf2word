import fitz
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import logging
from datetime import datetime
from flask import current_app
import numpy as np
import cv2
import tempfile
from PIL import Image
import io
import re
from collections import defaultdict
from .analysis.pattern_matcher import PatternMatcher
import subprocess
import shutil
from statistics import StatisticsError, mode, mean
import json

# Add new imports for the hybrid approach
try:
    from pdf2docx import Converter as Pdf2DocxConverter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False
    
try:
    import camelot
    CAMELOT_AVAILABLE = False
except ImportError:
    CAMELOT_AVAILABLE = False

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

class DocumentConverter:
    def __init__(self):
        self.color_scheme = None
        self.shape_patterns = None
        self.section_styles = {}
        self.detected_fonts = {}
        self.design_elements = {
            'main_color': None,
            'accent_color': None,
            'header_style': None,
            'has_borders': False,
            'has_background_elements': False
        }
        
        # Add cleanup of old files
        self._cleanup_old_files()

    def _cleanup_old_files(self, max_age_hours=24):
        """Clean up old temporary files"""
        try:
            temp_dirs = [
                os.path.join(current_app.root_path, 'temp', 'uploads'),
                os.path.join(current_app.root_path, 'temp', 'converted')
            ]
            
            current_time = datetime.now()
            
            for temp_dir in temp_dirs:
                if not os.path.exists(temp_dir):
                    continue
                    
                for filename in os.listdir(temp_dir):
                    filepath = os.path.join(temp_dir, filename)
                    file_modified = datetime.fromtimestamp(os.path.getmtime(filepath))
                    
                    if (current_time - file_modified).total_seconds() > max_age_hours * 3600:
                        os.remove(filepath)
                        
        except Exception as e:
            logger.error(f"Error cleaning up temporary files: {str(e)}")

    def convert(self, input_path, output_path, target_format=None):
        """Convert document to target format using the best available method"""
        try:
            logger.debug(f"Starting conversion from {input_path} to {output_path}")
            
            # Basic validation
            if not os.path.exists(input_path):
                raise FileNotFoundError(f"Input file not found: {input_path}")
            
            # Create output directory if it doesn't exist
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Perform conversion based on target format
            if target_format is None or target_format.lower() == 'docx':
                # Use the new hybrid approach for DOCX conversion
                return self.hybrid_convert_to_docx(input_path, output_path)
            elif target_format.lower() == 'txt':
                return self.convert_to_txt(input_path, output_path)
            else:
                raise ValueError(f"Unsupported target format: {target_format}")
            
        except Exception as e:
            logger.error(f"Conversion error: {str(e)}")
            raise
            
    def hybrid_convert_to_docx(self, input_path, output_path):
        """Hybrid approach for PDF to DOCX conversion using multiple engines"""
        try:
            # Step 1: Analyze document to determine type and complexity
            doc_type, doc_complexity = self._analyze_document_type(input_path)
            logger.debug(f"Detected document type: {doc_type}, complexity: {doc_complexity}")
            
            # Step 2: Choose the best conversion engine based on document type
            if doc_type == 'resume' and PDF2DOCX_AVAILABLE:
                # Resumes typically convert better with pdf2docx
                logger.debug("Using pdf2docx engine for resume conversion")
                result = self._convert_with_pdf2docx(input_path, output_path)
            elif doc_type == 'table_heavy' and CAMELOT_AVAILABLE:
                # Documents with many tables may benefit from camelot + our custom processing
                logger.debug("Using camelot-enhanced conversion for table-heavy document")
                result = self._convert_with_camelot_enhanced(input_path, output_path)
            elif doc_complexity == 'complex':
                # Try multiple engines and select the best result
                logger.debug("Complex document detected, trying multiple engines")
                result = self._convert_with_multiple_engines(input_path, output_path)
            else:
                # Use our standard conversion for simple documents
                logger.debug("Using standard conversion engine")
                result = self.convert_to_docx(input_path, output_path)
            
            # Step 3: Apply specialized post-processing based on document type
            self._apply_specialized_post_processing(output_path, doc_type)
            
            return result
            
        except Exception as e:
            logger.error(f"Hybrid conversion error: {str(e)}")
            # Fall back to standard conversion if hybrid approach fails
            logger.debug("Falling back to standard conversion")
            return self.convert_to_docx(input_path, output_path)
            
    def _analyze_document_type(self, input_path):
        """Analyze document to determine its type and complexity"""
        try:
            pdf = fitz.open(input_path)
            page_count = pdf.page_count
            
            # Initialize counters
            total_words = 0
            resume_score = 0
            table_score = 0
            form_score = 0
            complexity_score = 0
            
            # Check first 3 pages max for efficiency
            for page_num in range(min(3, page_count)):
                page = pdf[page_num]
                
                # Get page text
                text = page.get_text()
                words = re.findall(r'\w+', text)
                total_words += len(words)
                
                # Resume detection keywords
                resume_keywords = ['resume', 'cv', 'curriculum vitae', 'professional experience', 
                                  'skills', 'education', 'work history', 'profile', 'objective',
                                  'summary', 'utbildning', 'arbetslivserfarenhet']
                
                # Count resume keywords
                for keyword in resume_keywords:
                    if keyword.lower() in text.lower():
                        resume_score += 1
                
                # Check for email/phone patterns (common in resumes)
                if re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text):  # Email
                    resume_score += 2
                if re.search(r'(?:\+|00)?[0-9()\s-]{7,}', text):  # Phone
                    resume_score += 1
                
                # Table detection
                rect_count = 0
                for drawing in page.get_drawings():
                    if drawing.get("type") == "rect":
                        rect_count += 1
                
                # High number of rectangles often indicates tables or forms
                if rect_count > 10:
                    table_score += rect_count // 5
                    form_score += rect_count // 10
                
                # Complex layout detection
                blocks = page.get_text("dict").get("blocks", [])
                if len(blocks) > 20:  # Many text blocks suggests complex layout
                    complexity_score += len(blocks) // 10
                
                # Check for multi-column layout (increases complexity)
                columns = self._detect_columns(page)
                if len(columns) > 1:
                    complexity_score += len(columns) * 2
            
            # Determine document type based on scores
            doc_type = 'general'
            if resume_score >= 5:
                doc_type = 'resume'
            elif table_score >= 10:
                doc_type = 'table_heavy'
            elif form_score >= 8:
                doc_type = 'form'
            
            # Determine complexity
            complexity = 'simple'
            if complexity_score >= 10 or page_count > 20:
                complexity = 'complex'
            elif complexity_score >= 5:
                complexity = 'moderate'
            
            pdf.close()
            return doc_type, complexity
            
        except Exception as e:
            logger.warning(f"Error analyzing document type: {str(e)}")
            return 'general', 'moderate'  # Default values if analysis fails
            
    def _detect_columns(self, page):
        """Simple column detection for document analysis"""
        try:
            # Get page dimensions
            page_width = page.rect.width
            
            # Get text blocks
            blocks = page.get_text("dict").get("blocks", [])
            x_coordinates = []
            
            # Collect x-coordinates of all text blocks
            for block in blocks:
                if block.get("type") == 0 and "bbox" in block:
                    x_min = block["bbox"][0]
                    x_max = block["bbox"][2]
                    x_coordinates.extend([x_min, x_max])
            
            if not x_coordinates:
                return []
                
            # Create histogram to identify columns
            hist, bins = np.histogram(x_coordinates, bins=20, range=(0, page_width))
            
            # Find significant valleys in the histogram (gaps between columns)
            valleys = []
            for i in range(1, len(hist) - 1):
                if (hist[i] < hist[i-1] and hist[i] < hist[i+1] and 
                    hist[i] < max(hist) * 0.2):  # Significant gap
                    valleys.append(bins[i])
            
            # Create columns from valleys
            columns = []
            if valleys:
                # Add left edge as first boundary
                boundaries = [0] + valleys + [page_width]
                
                for i in range(len(boundaries) - 1):
                    left = boundaries[i]
                    right = boundaries[i+1]
                    width_ratio = (right - left) / page_width
                    
                    # Skip if column is too narrow
                    if width_ratio < 0.1:
                        continue
                    
                    columns.append({
                        'left': left,
                        'right': right,
                        'width_ratio': width_ratio
                    })
            
            return columns
            
        except Exception as e:
            logger.debug(f"Error in simple column detection: {str(e)}")
            return []
    
    def _convert_with_pdf2docx(self, input_path, output_path):
        """Convert PDF to DOCX using pdf2docx library"""
        try:
            if not PDF2DOCX_AVAILABLE:
                logger.warning("pdf2docx library not available, falling back to standard conversion")
                return self.convert_to_docx(input_path, output_path)
            
            # Use pdf2docx for conversion
            cv = Pdf2DocxConverter(input_path)
            cv.convert(output_path)
            cv.close()
            
            # Apply additional post-processing specific to pdf2docx output
            self._post_process_pdf2docx_output(output_path)
            
            return True
            
        except Exception as e:
            logger.error(f"pdf2docx conversion error: {str(e)}")
            # Fall back to standard conversion
            logger.debug("Falling back to standard conversion")
            return self.convert_to_docx(input_path, output_path)
    
    def _post_process_pdf2docx_output(self, docx_path):
        """Apply post-processing to fix common issues in pdf2docx output"""
        try:
            # Load the document
            doc = Document(docx_path)
            
            # Fix 1: Handle empty table cells better
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Remove empty paragraphs in cells
                        for para in list(cell.paragraphs):
                            if not para.text.strip():
                                p = para._element
                                if p.getparent() is not None:
                                    p.getparent().remove(p)
            
            # Fix 2: Remove borders from all tables
            for table in doc.tables:
                self._remove_table_borders_completely(table)
            
            # Fix 3: Convert drawing tables to proper tables if needed
            # (This is a complex process, simplified here)
            
            # Fix 4: Apply consistent paragraph spacing
            for para in doc.paragraphs:
                if not para.style.name.startswith('Heading'):
                    para.paragraph_format.space_after = Pt(6)
            
            # Save the document
            doc.save(docx_path)
            
        except Exception as e:
            logger.warning(f"Error in pdf2docx post-processing: {str(e)}")
    
    def _convert_with_camelot_enhanced(self, input_path, output_path):
        """Convert PDF with enhanced table processing using camelot"""
        try:
            if not CAMELOT_AVAILABLE:
                logger.warning("Camelot library not available, falling back to standard conversion")
                return self.convert_to_docx(input_path, output_path)
            
            # Use our standard conversion first
            self.convert_to_docx(input_path, output_path)
            
            # Extract tables using camelot
            tables = camelot.read_pdf(input_path)
            
            # Only proceed if we found tables
            if len(tables) > 0:
                # Load the document
                doc = Document(output_path)
                
                # For each table found by camelot
                for i, table in enumerate(tables):
                    # Convert camelot table to pandas DataFrame
                    df = table.df
                    
                    # Create a Word table
                    rows, cols = df.shape
                    word_table = doc.add_table(rows=rows, cols=cols)
                    word_table.style = 'Table Grid'
                    
                    # Fill the table with data
                    for r in range(rows):
                        for c in range(cols):
                            word_table.cell(r, c).text = df.iloc[r, c]
                    
                    # Add page break after each table except the last one
                    if i < len(tables) - 1:
                        doc.add_page_break()
                
                # Remove any empty tables from the original conversion
                self._clean_empty_tables(doc)
                
                # Save the document
                doc.save(output_path)
            
            return True
            
        except Exception as e:
            logger.error(f"Camelot-enhanced conversion error: {str(e)}")
            # Document already converted with standard method
            return True
    
    def _clean_empty_tables(self, doc):
        """Remove empty tables from document"""
        for table in list(doc.tables):
            is_empty = True
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        is_empty = False
                        break
                if not is_empty:
                    break
            
            if is_empty:
                # Remove the table
                tbl = table._tbl
                parent = tbl.getparent()
                if parent is not None:
                    parent.remove(tbl)
    
    def _convert_with_multiple_engines(self, input_path, output_path):
        """Try multiple conversion engines and select the best result"""
        try:
            # Create temporary files for each approach
            temp_dir = tempfile.mkdtemp()
            standard_output = os.path.join(temp_dir, "standard.docx")
            pdf2docx_output = os.path.join(temp_dir, "pdf2docx.docx")
            
            # Create base filename for output path
            base_output = os.path.splitext(output_path)[0]
            
            # 1. Standard conversion
            logger.debug("Trying standard conversion")
            standard_success = self.convert_to_docx(input_path, standard_output)
            
            # 2. pdf2docx conversion if available
            pdf2docx_success = False
            if PDF2DOCX_AVAILABLE:
                logger.debug("Trying pdf2docx conversion")
                try:
                    pdf2docx_success = self._convert_with_pdf2docx(input_path, pdf2docx_output)
                except Exception as e:
                    logger.warning(f"pdf2docx conversion failed: {str(e)}")
            
            # Evaluate results to select the best approach
            standard_score = 0
            pdf2docx_score = 0
            
            if standard_success:
                standard_score = self._evaluate_conversion_quality(standard_output)
                logger.debug(f"Standard conversion quality score: {standard_score}")
            
            if pdf2docx_success:
                pdf2docx_score = self._evaluate_conversion_quality(pdf2docx_output)
                logger.debug(f"pdf2docx conversion quality score: {pdf2docx_score}")
            
            # Select best result based on scores
            if pdf2docx_success and pdf2docx_score > standard_score:
                logger.debug("Using pdf2docx result (higher quality)")
                shutil.copy(pdf2docx_output, output_path)
            elif standard_success:
                logger.debug("Using standard conversion result")
                shutil.copy(standard_output, output_path)
            else:
                logger.warning("All conversion attempts failed, using fallback")
                # Create minimal document as fallback
                doc = Document()
                doc.add_paragraph("Conversion failed. Please try a different format.")
                doc.save(output_path)
            
            # Clean up temp directory
            try:
                shutil.rmtree(temp_dir)
            except Exception as e:
                logger.debug(f"Error cleaning up temp directory: {str(e)}")
            
            return True
            
        except Exception as e:
            logger.error(f"Multiple engine conversion error: {str(e)}")
            # Fall back to standard conversion
            return self.convert_to_docx(input_path, output_path)
    
    def _evaluate_conversion_quality(self, docx_path):
        """Evaluate the quality of a conversion result"""
        try:
            # Load the document
            doc = Document(docx_path)
            
            # Initialize score
            score = 0
            
            # Feature 1: Text extraction completeness
            text_length = 0
            for para in doc.paragraphs:
                text_length += len(para.text)
            
            # Longer text generally means more complete extraction
            if text_length > 1000:
                score += 10
            elif text_length > 500:
                score += 5
            elif text_length > 100:
                score += 2
            
            # Feature 2: Structural elements preservation
            headings_count = 0
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    headings_count += 1
            
            # Documents with proper heading structure get bonus points
            score += min(10, headings_count * 2)
            
            # Feature 3: Table quality
            tables_count = len(doc.tables)
            tables_score = 0
            
            for table in doc.tables:
                # Check if table has content
                has_content = False
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            has_content = True
                            break
                    if has_content:
                        break
                
                if has_content:
                    tables_score += 2
                    
                    # Bonus for larger tables (more complex)
                    if len(table.rows) > 5 and len(table.columns) > 3:
                        tables_score += 3
            
            score += tables_score
            
            # Feature 4: Image preservation
            # Count inline shapes (images)
            image_count = 0
            for shape in doc.inline_shapes:
                image_count += 1
            
            score += min(10, image_count * 2)
            
            # Feature 5: Formatting preservation (simplified)
            formatting_score = 0
            for para in doc.paragraphs:
                # Check for mixed formatting in paragraph (indicates rich formatting preservation)
                formats_in_para = set()
                for run in para.runs:
                    format_key = (run.bold, run.italic, run.underline)
                    formats_in_para.add(format_key)
                
                if len(formats_in_para) > 1:
                    formatting_score += 1
            
            score += min(10, formatting_score)
            
            return score
            
        except Exception as e:
            logger.warning(f"Error evaluating conversion quality: {str(e)}")
            return 0
    
    def _apply_specialized_post_processing(self, docx_path, doc_type):
        """Apply document-type-specific post-processing"""
        try:
            doc = Document(docx_path)
            
            # Common fixes for all document types
            self._fix_character_spacing(doc)
            self._fix_table_borders(doc)
            self._cleanup_empty_paragraphs(doc)
            
            # Type-specific processing
            if doc_type == 'resume':
                self._enhance_resume_formatting(doc)
            elif doc_type == 'table_heavy':
                self._optimize_table_layout(doc)
            elif doc_type == 'form':
                self._preserve_form_layout(doc)
            
            # Save the document
            doc.save(docx_path)
            
        except Exception as e:
            logger.warning(f"Error in specialized post-processing: {str(e)}")
    
    def _fix_character_spacing(self, doc):
        """Fix character spacing issues in the document"""
        for para in doc.paragraphs:
            # Look for text with unusual spacing patterns
            text = para.text
            
            # Fix 1: Remove excessive spaces
            if '  ' in text:
                new_text = re.sub(r' {2,}', ' ', text)
                
                # Only update if changed
                if new_text != text:
                    # Clear the paragraph and add the fixed text
                    for run in list(para.runs):
                        p = run._element
                        p.getparent().remove(p)
                    
                    para.add_run(new_text)
    
    def _fix_table_borders(self, doc):
        """Fix table borders throughout the document"""
        for table in doc.tables:
            # Remove all borders
            self._remove_table_borders_completely(table)
            
            # Optionally add minimal styling if needed
            table.style = 'Table Grid'
    
    def _cleanup_empty_paragraphs(self, doc):
        """Remove excessive empty paragraphs"""
        # Find sequences of empty paragraphs
        empty_sequences = []
        current_sequence = []
        
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                current_sequence.append(i)
            else:
                if len(current_sequence) > 1:  # Keep single empty paragraphs for spacing
                    empty_sequences.append(current_sequence)
                current_sequence = []
        
        # Add the last sequence if it exists
        if len(current_sequence) > 1:
            empty_sequences.append(current_sequence)
        
        # Remove all but the first paragraph in each sequence
        paragraphs_to_remove = []
        for sequence in empty_sequences:
            # Keep the first empty paragraph for spacing
            paragraphs_to_remove.extend(sequence[1:])
        
        # Remove paragraphs in reverse order to avoid index issues
        for idx in sorted(paragraphs_to_remove, reverse=True):
            p = doc.paragraphs[idx]._element
            p.getparent().remove(p)
    
    def _enhance_resume_formatting(self, doc):
        """Enhance formatting specifically for resumes"""
        # Identify common resume sections
        resume_sections = ['summary', 'profile', 'experience', 'education', 
                           'skills', 'certifications', 'languages',
                           'arbetslivserfarenhet', 'utbildning', 'färdigheter']
        
        for para in doc.paragraphs:
            text = para.text.lower().strip()
            
            # Look for section headers
            is_section = False
            for section in resume_sections:
                if text.startswith(section) or text == section:
                    is_section = True
                    break
            
            if is_section:
                # Format as section header
                para.style = 'Heading 2'
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(6)
                
                # Make the text bold
                for run in para.runs:
                    run.bold = True
    
    def _optimize_table_layout(self, doc):
        """Optimize table layouts in table-heavy documents"""
        for table in doc.tables:
            # Set consistent cell margins
            for row in table.rows:
                for cell in row.cells:
                    # Access cell properties
                    try:
                        tc = cell._element.tcPr
                        if tc is None:
                            tc = OxmlElement('w:tcPr')
                            cell._element.append(tc)
                            
                        # Add cell margins specification
                        tcMar = OxmlElement('w:tcMar')
                        
                        # Set all margins (1pt = 20 dxa)
                        for side in ['top', 'left', 'bottom', 'right']:
                            margin = OxmlElement(f'w:{side}')
                            margin.set(qn('w:w'), '60')  # 3pt margin
                            margin.set(qn('w:type'), 'dxa')
                            tcMar.append(margin)
                            
                        # Check if margins already exist
                        existing_tcMar = tc.find('.//w:tcMar')
                        if existing_tcMar is not None:
                            tc.remove(existing_tcMar)
                            
                        tc.append(tcMar)
                    except Exception as e:
                        logger.debug(f"Error setting cell margins: {str(e)}")
            
            # Fix header row if present
            if len(table.rows) > 0:
                # Assume first row could be header
                header_row = table.rows[0]
                
                # Look for header-like formatting
                header_like = True
                for cell in header_row.cells:
                    if not any(run.bold for p in cell.paragraphs for run in p.runs):
                        header_like = False
                        break
                
                if header_like:
                    # Enhance header formatting
                    for cell in header_row.cells:
                        for para in cell.paragraphs:
                            para.paragraph_format.space_after = Pt(2)
                            for run in para.runs:
                                run.bold = True
    
    def _preserve_form_layout(self, doc):
        """Preserve form layout elements"""
        # Forms often have specific layout needs
        # 1. Look for form fields (text followed by lines/underscores or blank spaces)
        for para in doc.paragraphs:
            text = para.text
            
            # Check for form field patterns
            if re.search(r'[^:]+:\s*_{3,}', text) or re.search(r'[^:]+:(\s{3,}|\t)', text):
                # This might be a form field label
                parts = re.split(r'(:|\s{3,}|_{3,})', text, 1)
                if len(parts) >= 2:
                    label = parts[0].strip()
                    
                    # Clear the paragraph
                    for run in list(para.runs):
                        p = run._element
                        p.getparent().remove(p)
                    
                    # Add formatted label and field
                    label_run = para.add_run(label + ": ")
                    label_run.bold = True
                    
                    # Add a tab
                    para.paragraph_format.tab_stops.add_tab_stop(Inches(2.5))
                    para.add_run("\t")
                    
                    # Add field placeholder (can be empty)
                    if len(parts) > 2:
                        field_value = parts[2].strip()
                        para.add_run(field_value)

    def convert_to_txt(self, input_path, output_path):
        """Convert PDF to plain text"""
        pdf = None
        try:
            # Open PDF
            pdf = fitz.open(input_path)
            
            # Extract text from all pages
            text = ""
            for page_num in range(pdf.page_count):
                page = pdf[page_num]
                text += page.get_text()
                text += "\n\n"  # Add page separators
            
            # Write to output file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
                
            return True
            
        except Exception as e:
            logger.error(f"PDF to TXT conversion error: {str(e)}")
            raise
            
        finally:
            if pdf:
                pdf.close()

    def convert_to_docx(self, input_path, output_path):
        """Convert PDF to DOCX with layout preservation"""
        pdf = None
        try:
            # Create Word document
            doc = Document()
            
            # Set minimal default margins for better layout
            for section in doc.sections:
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
            
            # Open PDF
            pdf = fitz.open(input_path)
            logger.debug(f"Opened PDF with {pdf.page_count} pages")
            
            # NEW: Check first page for decorative elements at top
            if pdf.page_count > 0:
                self._has_decorative_header = self._check_for_decorative_header(pdf[0])
                logger.debug(f"Decorative header detected: {self._has_decorative_header}")
            else:
                self._has_decorative_header = False
            
            # Analyze whole document to detect global layout patterns
            layout_info_all_pages = []
            for page_num in range(min(pdf.page_count, 3)):  # Limit to first 3 pages
                page = pdf[page_num]
                layout_info = self._analyze_page_layout(page)
                layout_info_all_pages.append(layout_info)
            
            # Determine global document style based on all analyzed pages
            global_layout = self._combine_layout_info(layout_info_all_pages)
            
            # Set consistent document styles based on analysis
            self._set_document_styles(doc, global_layout)
            
            # Process each page
            for page_num in range(pdf.page_count):
                page = pdf[page_num]
                
                # Only add page break after first page
                if page_num > 0:
                    doc.add_page_break()
                    # Ensure there's a section for this page
                    if len(doc.sections) <= page_num:
                        doc.add_section()
                
                # Set page size and orientation
                # Fix potential index error by making sure the section exists
                if page_num < len(doc.sections):
                    section = doc.sections[page_num]
                    self._set_page_properties(section, page)
                else:
                    # If section doesn't exist, use the last section
                    section = doc.sections[-1]
                    self._set_page_properties(section, page)
                
                # Extract layout information
                layout_info = self._analyze_page_layout(page)
                
                # Skip decorative headers on first page if detected
                if page_num == 0 and self._has_decorative_header:
                    layout_info['skip_decorative_top'] = True
                
                # Use global layout type if more consistent
                if global_layout.get('consistent_layout_type'):
                    layout_info['type'] = global_layout['layout_type']
                    
                logger.debug(f"Detected layout type: {layout_info['type']}")
                
                # Process page based on layout type
                if layout_info['type'] == 'multi_column' and layout_info['columns']:
                    self._process_multi_column_page(doc, page, layout_info)
                else:
                    self._process_single_column_page(doc, page, layout_info)
            
            # Post-process the document for final cleanup and adjustments
            self._post_process_document(doc)
            
            # Save document
            doc.save(output_path)
            logger.debug(f"Saved document to {output_path}")
            
            return True
            
        except Exception as e:
            logger.error(f"PDF to DOCX conversion error: {str(e)}")
            raise
            
        finally:
            if pdf:
                pdf.close()

    def _check_for_decorative_header(self, page):
        """Check if the page has decorative elements at the top"""
        try:
            # Get page size
            page_width = page.rect.width
            page_height = page.rect.height
            
            # Define top region (top 15% of page)
            top_region_height = page_height * 0.15
            
            # Get text blocks in the top region
            blocks = page.get_text("dict").get("blocks", [])
            top_blocks = []
            
            for block in blocks:
                if "bbox" in block and len(block["bbox"]) >= 4:
                    y_top = block["bbox"][1]
                    if y_top < top_region_height:
                        top_blocks.append(block)
            
            # No blocks at top - less likely to have decorative elements
            if not top_blocks:
                return False
                
            # Check for any lines or rectangles in the top region
            # These might be decorative elements
            paths = page.get_drawings()
            top_paths = []
            
            for path in paths:
                # Check if path intersects with top region
                if path["rect"][1] < top_region_height:
                    top_paths.append(path)
            
            # If we have paths in the top region, might be decorative
            if top_paths:
                return True
                
            # Check for very short lines of text (often decorative)
            short_text_count = 0
            for block in top_blocks:
                if block.get("type") == 0:  # Text block
                    text = "".join(span.get("text", "") for line in block.get("lines", []) 
                                  for span in line.get("spans", []))
                    if len(text.strip()) < 3:
                        short_text_count += 1
            
            # If more than 2 very short text blocks at top, might be decorative
            if short_text_count > 1:
                return True
                
            # Check for horizontal lines spanning most of the page width
            for path in top_paths:
                # If path is a horizontal line
                if (abs(path["rect"][3] - path["rect"][1]) < 2 and  # Height ≈ 0
                    (path["rect"][2] - path["rect"][0]) > page_width * 0.7):  # Width > 70% of page
                    return True
            
            return False
            
        except Exception as e:
            logger.debug(f"Error checking for decorative header: {str(e)}")
            return False
            
    def _process_single_column_page(self, doc, page, layout_info=None):
        """Process a page with single-column layout"""
        try:
            # If layout_info wasn't provided, create a default
            if layout_info is None:
                layout_info = {
                    'type': 'single_column',
                    'skip_decorative_top': False
                }
                
            # Extract text blocks
            blocks = page.get_text("dict").get("blocks", [])
            
            # Sort blocks by y-position (top to bottom)
            text_blocks = [b for b in blocks if b.get("type") == 0 and "bbox" in b]
            text_blocks.sort(key=lambda b: b["bbox"][1])
            
            # Skip decorative elements at top if detected
            if layout_info.get('skip_decorative_top', False) and text_blocks:
                # Skip first block if it's in the top 10% of page and has little text
                if text_blocks[0]["bbox"][1] < page.rect.height * 0.1:
                    first_text = self._extract_text(text_blocks[0]).strip()
                    if len(first_text) < 10:  # Very short text, likely decorative
                        text_blocks = text_blocks[1:]
            
            # Process each block
            for block in text_blocks:
                # Extract text and formatting
                paragraph = doc.add_paragraph()
                
                for line_idx, line in enumerate(block.get("lines", [])):
                    if line_idx > 0:
                        # Add line break between lines
                        paragraph.add_run().add_break()
                    
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if not text:
                            continue
                        
                        run = paragraph.add_run(text + " ")
                        self._apply_span_formatting(run, span)
                
                # Check if this might be a header
                if any(span.get("size", 0) > 12 or (span.get("flags", 0) & 16) for line in block.get("lines", []) for span in line.get("spans", [])):
                    paragraph.style = "Heading 2"
                    
                    # Add some spacing
                    paragraph.paragraph_format.space_before = Pt(12)
                    paragraph.paragraph_format.space_after = Pt(6)
                else:
                    # Regular paragraph
                    paragraph.paragraph_format.space_after = Pt(6)
            
            # Process images if any
            self._extract_and_add_images(doc, page)
        except Exception as e:
            logger.warning(f"Error in single column processing: {str(e)}")
            # Create a simple paragraph with the page text as fallback
            try:
                doc.add_paragraph(page.get_text())
            except:
                pass
    
    def _post_process_document(self, doc):
        """Apply final adjustments to the document before saving"""
        try:
            # Fix empty paragraphs (excessive spacing)
            for para in doc.paragraphs:
                if not para.text.strip():
                    # Remove excessive empty paragraphs (keep some for spacing)
                    if hasattr(para, '_p') and hasattr(para._p, 'getparent'):
                        parent = para._p.getparent()
                        if parent is not None:
                            # Check if previous paragraph was also empty
                            prev_was_empty = False
                            prev_idx = list(parent).index(para._p) - 1
                            if prev_idx >= 0:
                                prev_p = parent[prev_idx]
                                if hasattr(prev_p, 'text') and not prev_p.text.strip():
                                    prev_was_empty = True
                            
                            # Remove if previous was also empty (keep one for spacing)
                            if prev_was_empty:
                                parent.remove(para._p)
            
            # ENHANCED: Aggressive table border and empty table removal
            self._aggressively_clean_tables(doc)
                            
        except Exception as e:
            logger.warning(f"Error in post-processing document: {str(e)}")
            # Continue despite errors in post-processing

    def _aggressively_clean_tables(self, doc):
        """More aggressive table cleaning focusing on document start"""
        try:
            for i, table in enumerate(list(doc.tables)):  # Use list to allow safe removal during iteration
                # Check if table is empty or contains only whitespace
                table_is_empty = True
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            table_is_empty = True
                            # Do one more check - sometimes there's hidden content
                            has_real_content = False
                            for paragraph in cell.paragraphs:
                                if paragraph.text.strip():
                                    has_real_content = True
                                    break
                            if has_real_content:
                                table_is_empty = False
                                break
                    if not table_is_empty:
                        break
                
                # Special handling for document start tables (more likely to be decorative)
                is_document_start = (i == 0)
                
                if table_is_empty:
                    # Remove empty table completely by getting its parent and removing it
                    if hasattr(table, '_tbl') and hasattr(table._tbl, 'getparent'):
                        parent = table._tbl.getparent()
                        if parent is not None:
                            # Check if this is the first table in a fresh document
                            # If so, we want to be extra aggressive in removal
                            try:
                                parent.remove(table._tbl)
                                logger.debug(f"Successfully removed empty table #{i}")
                            except Exception as e:
                                logger.warning(f"Error removing table: {str(e)}")
                else:
                    # Apply stronger border removal for all tables
                    self._remove_table_borders_completely(table, is_document_start)
        except Exception as e:
            logger.warning(f"Error in aggressive table cleaning: {str(e)}")

    def _remove_table_borders_completely(self, table, is_document_start=False):
        """Complete border removal with multiple fallback methods"""
        try:
            # Step 1: Try enhanced border removal
            self._remove_table_borders_enhanced(table)
            
            # Step 2: Direct XML approach to absolutely ensure borders are gone
            tbl = table._tbl
            
            # Add explicit border-related namespaces
            nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            # Define all possible border elements that could appear
            border_paths = [
                './/w:tblBorders',
                './/w:tcBorders',
                './/w:tblBorders/w:top',
                './/w:tblBorders/w:left',
                './/w:tblBorders/w:bottom',
                './/w:tblBorders/w:right',
                './/w:tblBorders/w:insideH',
                './/w:tblBorders/w:insideV',
                './/w:pBdr'  # Paragraph borders that might be in cells
            ]
            
            # Remove each type of border element
            for path in border_paths:
                elements = tbl.xpath(path, namespaces=nsmap)
                for element in elements:
                    parent = element.getparent()
                    if parent is not None:
                        parent.remove(element)
            
            # Step 3: Explicitly set zero-width border on table properties
            try:
                tblPr = tbl.find('.//w:tblPr', namespaces=nsmap)
                if tblPr is not None:
                    # Create new border settings with nil values
                    new_borders = OxmlElement('w:tblBorders')
                    
                    for border_type in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                        border = OxmlElement(f'w:{border_type}')
                        border.set(qn('w:val'), 'nil')
                        border.set(qn('w:sz'), '0')
                        border.set(qn('w:space'), '0')
                        border.set(qn('w:color'), 'auto')
                        new_borders.append(border)
                    
                    # Add borders element to table properties
                    existing = tblPr.find('.//w:tblBorders', namespaces=nsmap)
                    if existing is not None:
                        tblPr.remove(existing)
                    tblPr.append(new_borders)
            except Exception as e:
                logger.debug(f"Error setting table properties: {str(e)}")
            
            # Step 4: Handle cells individually
            for row in table.rows:
                for cell in row.cells:
                    try:
                        # Handle each cell's borders explicitly
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        
                        # Remove any existing borders
                        tcBorders = tcPr.find('.//w:tcBorders', namespaces=nsmap)
                        if tcBorders is not None:
                            tcPr.remove(tcBorders)
                        
                        # Add new borders with "nil" value
                        tcBorders = OxmlElement('w:tcBorders')
                        
                        for border in ['top', 'left', 'bottom', 'right']:
                            border_elem = OxmlElement(f'w:{border}')
                            border_elem.set(qn('w:val'), 'nil')  # nil = no border
                            border_elem.set(qn('w:sz'), '0')     # zero width
                            border_elem.set(qn('w:space'), '0')  # no spacing
                            border_elem.set(qn('w:color'), 'auto')
                            tcBorders.append(border_elem)
                            
                        tcPr.append(tcBorders)
                        
                        # Special handling for first row top borders in document-start tables
                        if is_document_start and row._index == 0:
                            # Extra strong removal for top row border which tends to persist
                            for paragraph in cell.paragraphs:
                                if hasattr(paragraph, '_p'):
                                    p = paragraph._p
                                    pPr = p.get_or_add_pPr()
                                    pBdr = pPr.find('.//w:pBdr')
                                    if pBdr is not None:
                                        pPr.remove(pBdr)
                    except Exception as e:
                        logger.debug(f"Error processing cell: {str(e)}")
            
            # Step 5: Set table to look like it has no borders in all attributes
            table.style = 'Table Grid'  # Reset to standard style
            
            # Remove any styling that might add borders
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.style and paragraph.style.name != 'Normal':
                            paragraph.style = 'Normal'
                            
        except Exception as e:
            logger.warning(f"Enhanced border removal failed: {str(e)}")
            # Try standard removal as last resort
            self._remove_table_borders(table)
    
    def _set_document_styles(self, doc, layout_info):
        """Apply consistent styles throughout the document based on analysis"""
        try:
            # Get the style information
            header_sizes = layout_info.get('header_font_sizes', [16, 14, 12])
            body_size = layout_info.get('body_font_size', 11)
            base_font = layout_info.get('base_font', 'Calibri')
            
            # Set default styles to match original document
            styles = doc.styles
            
            # Body text style
            if 'Normal' in styles:
                style = styles['Normal']
                style.font.name = base_font
                style.font.size = Pt(body_size)
                style.paragraph_format.space_after = Pt(6)
            
            # Heading styles
            if header_sizes and len(header_sizes) >= 1 and 'Heading 1' in styles:
                style = styles['Heading 1']
                style.font.name = base_font
                style.font.size = Pt(header_sizes[0])
                style.font.bold = True
                style.paragraph_format.space_before = Pt(12)
                style.paragraph_format.space_after = Pt(6)
            
            if header_sizes and len(header_sizes) >= 2 and 'Heading 2' in styles:
                style = styles['Heading 2']
                style.font.name = base_font
                style.font.size = Pt(header_sizes[min(1, len(header_sizes)-1)])
                style.font.bold = True
                style.paragraph_format.space_before = Pt(10)
                style.paragraph_format.space_after = Pt(6)
            
            if header_sizes and len(header_sizes) >= 3 and 'Heading 3' in styles:
                style = styles['Heading 3']
                style.font.name = base_font
                style.font.size = Pt(header_sizes[min(2, len(header_sizes)-1)])
                style.font.bold = True
                style.paragraph_format.space_before = Pt(8)
                style.paragraph_format.space_after = Pt(4)
            
            # List style
            if 'List Bullet' in styles:
                style = styles['List Bullet']
                style.font.name = base_font
                style.font.size = Pt(body_size)
                style.paragraph_format.left_indent = Inches(0.25)
                style.paragraph_format.space_after = Pt(3)
            
        except Exception as e:
            logger.warning(f"Error setting document styles: {str(e)}")

    def _process_multi_column_page(self, doc, page, layout_info):
        """Process a page with multi-column layout"""
        try:
            # Add a table for multi-column layout
            num_columns = len(layout_info['columns'])
            if num_columns < 1:
                # Fall back to single column if no columns detected
                self._process_single_column_page(doc, page)
                return
                
            table = doc.add_table(rows=1, cols=num_columns)
            table.style = 'Table Grid'
            table.autofit = False
            
            # Remove table borders
            self._remove_table_borders(table)
            
            # Calculate optimal column width ratios
            column_width_ratios = []
            for col in layout_info['columns']:
                width_ratio = col['width_ratio']
                # Ensure minimum width
                width_ratio = max(0.1, min(0.9, width_ratio))
                column_width_ratios.append(width_ratio)
                
            # Normalize ratios to sum to 1
            total_ratio = sum(column_width_ratios)
            if total_ratio > 0:
                normalized_ratios = [r/total_ratio for r in column_width_ratios]
            else:
                # Equal widths if ratios are invalid
                normalized_ratios = [1.0/num_columns] * num_columns
            
            # Apply column widths
            for i, width_ratio in enumerate(normalized_ratios):
                if i < len(table.columns):  # Ensure column exists
                    # Apply width based on page size and margins
                    effective_page_width = 6.5  # Assuming 8.5" page width with 1" total margins
                    column_width = width_ratio * effective_page_width
                    table.columns[i].width = Inches(column_width)
            
            # Extract and categorize blocks by column
            columns_content = self._categorize_blocks_by_column(page, layout_info['columns'])
            
            # Process each column
            for col_idx, blocks in enumerate(columns_content):
                if col_idx < len(table.cells):  # Ensure cell exists
                    cell = table.cell(0, col_idx)
                    
                    # Set cell margin for better spacing (3pt on all sides)
                    # Access cell's properties to set margins
                    try:
                        tc = cell._element.tcPr
                        if tc is None:
                            tc = OxmlElement('w:tcPr')
                            cell._element.append(tc)
                            
                        # Create margin specification
                        tcMar = OxmlElement('w:tcMar')
                        
                        # Set all margins (1pt = 20 dxa)
                        for side in ['top', 'left', 'bottom', 'right']:
                            margin = OxmlElement(f'w:{side}')
                            margin.set(qn('w:w'), '60')  # 3pt margin
                            margin.set(qn('w:type'), 'dxa')
                            tcMar.append(margin)
                            
                        tc.append(tcMar)
                    except Exception as e:
                        logger.debug(f"Error setting cell margins: {str(e)}")
                    
                    # Process blocks in this column
                    self._process_column_blocks(cell, blocks, layout_info)
                    
            # Set spacing between rows using paragraph spacing
            for i, row in enumerate(table.rows):
                for cell in row.cells:
                    if cell.paragraphs:
                        first_para = cell.paragraphs[0]
                        first_para.paragraph_format.space_before = Pt(0)
                        
                        # For all paragraphs
                        for para in cell.paragraphs:
                            para.paragraph_format.space_after = Pt(4)
        except Exception as e:
            logger.warning(f"Error in multi-column processing: {str(e)}")
            # Fall back to single column processing
            self._process_single_column_page(doc, page)
    
    def _remove_table_borders(self, table):
        """Remove borders from table"""
        try:
            # Access table properties to remove all borders
            tbl = table._tbl
            tblPr = tbl.get_or_add_tblPr()
            
            # Add table borders element
            tblBorders = OxmlElement('w:tblBorders')
            
            # Set all borders to none
            for border in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border_elem = OxmlElement(f'w:{border}')
                border_elem.set(qn('w:val'), 'none')
                border_elem.set(qn('w:sz'), '0')
                border_elem.set(qn('w:space'), '0')
                border_elem.set(qn('w:color'), 'auto')
                tblBorders.append(border_elem)
            
            # Replace any existing borders element
            existing_borders = tblPr.find('.//w:tblBorders')
            if existing_borders is not None:
                tblPr.remove(existing_borders)
            tblPr.append(tblBorders)
            
            # Also set individual cell borders to none
            for row in table.rows:
                for cell in row.cells:
                    # For all borders
                    for border in ['top', 'left', 'bottom', 'right']:
                        # Get the border element
                        border_elm = cell._element.xpath(f"w:tcPr/w:tcBorders/w:{border}")
                        if border_elm:
                            # Set the border value to 'nil'
                            border_elm[0].set(qn('w:val'), 'nil')
        except Exception as e:
            logger.warning(f"Error removing table borders: {str(e)}")
    
    def _categorize_blocks_by_column(self, page, columns):
        """Categorize blocks into columns based on position"""
        blocks = page.get_text("dict").get("blocks", [])
        columns_blocks = [[] for _ in columns]
        
        for block in blocks:
            if block.get("type") != 0 or "bbox" not in block:  # Skip non-text blocks or blocks without bbox
                continue
            
            # Get block coordinates
            x_min = block["bbox"][0]
            x_max = block["bbox"][2]
            block_center_x = (x_min + x_max) / 2
            
            # Find which column this block belongs to
            assigned = False
            
            # First try exact containment
            for col_idx, col_info in enumerate(columns):
                # Block is fully contained in column
                if col_info['left'] <= x_min and x_max <= col_info['right']:
                    columns_blocks[col_idx].append(block)
                    assigned = True
                    break
            
            # If not fully contained, try majority containment
            if not assigned:
                for col_idx, col_info in enumerate(columns):
                    # Block's center is in column
                    if col_info['left'] <= block_center_x <= col_info['right']:
                        columns_blocks[col_idx].append(block)
                        assigned = True
                        break
            
            # If block doesn't fit in any column, assign to nearest column
            if not assigned and columns:
                distances = [abs(block_center_x - (col['left'] + col['right'])/2) for col in columns]
                nearest_col = distances.index(min(distances))
                columns_blocks[nearest_col].append(block)
        
        # Sort blocks in each column by y-position
        for col_blocks in columns_blocks:
            col_blocks.sort(key=lambda b: b["bbox"][1])
        
        return columns_blocks
    
    def _process_column_blocks(self, cell, blocks, layout_info):
        """Process blocks within a column"""
        # Track the current position in the document
        current_section = None
        current_list = False
        
        # Check if this is likely a resume and use special processing
        is_resume = layout_info.get('is_resume', False)
        
        for i, block in enumerate(blocks):
            try:
                # Check if block is a header
                is_header = self._is_header_block(block, layout_info['headers'])
                
                # Create paragraph for this block
                if is_header:
                    # Add some spacing before headers (except first header)
                    if i > 0:
                        spacer = cell.add_paragraph()
                        spacer.paragraph_format.space_after = Pt(6)
                    
                    # Add header
                    header_para = cell.add_paragraph()
                    self._add_header_text(header_para, block, layout_info)
                    
                    # If it's a main header, update current section
                    section_name = self._extract_text(block).strip()
                    if section_name:
                        current_section = section_name
                        current_list = False
                else:
                    # Determine if this block might be a continuation of a previous block
                    is_continuation = False
                    if i > 0:
                        prev_block = blocks[i-1]
                        prev_text = self._extract_text(prev_block).strip()
                        current_text = self._extract_text(block).strip()
                        
                        # Check if previous block ends without punctuation and current starts lowercase
                        if (prev_text and not prev_text[-1] in '.!?:;' and 
                            current_text and current_text[0].islower()):
                            is_continuation = True
                    
                    # Special processing for continuation blocks
                    if is_continuation and cell.paragraphs:
                        # Add to the previous paragraph instead of creating a new one
                        para = cell.paragraphs[-1]
                        para.add_run(' ' + self._extract_text(block))
                    else:
                        # Process regular content
                        self._process_content_block(cell, block, layout_info, current_section, current_list, is_resume)
                    
                    # Check if this block contains bullet points
                    block_text = self._extract_text(block)
                    if any(bullet in block_text for bullet in layout_info['bullet_styles']):
                        current_list = True
                    elif block_text.strip() and not re.match(r'^\s*[•⦁◦·○●]\s', block_text):
                        current_list = False
            except Exception as e:
                logger.warning(f"Error processing block: {str(e)}")
                # Add a simple paragraph with the text to ensure content is not lost
                try:
                    p = cell.add_paragraph(self._extract_text(block))
                except:
                    pass
    
    def _is_header_block(self, block, headers):
        """Determine if a block is a header"""
        try:
            # Extract text from block
            block_text = self._extract_text(block).strip()
            
            # Check if any headers match this block
            for header in headers:
                if header['text'] in block_text:
                    return True
            
            # Check if block has header-like properties
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    # Check for larger font size or bold flag
                    if span.get("size", 0) > 13 or (span.get("flags", 0) & 16):
                        # Additionally check if text is short (typical for headers)
                        if len(block_text) < 50:
                            return True
        except Exception as e:
            logger.debug(f"Error checking header: {str(e)}")
        
        return False
    
    def _add_header_text(self, paragraph, block, layout_info):
        """Add properly formatted header text to paragraph"""
        try:
            text = self._extract_text(block).strip()
            if not text:
                return
            
            # Determine header level based on font size and position
            header_level = 1  # Default to level 1
            max_size = 0
            is_bold = False
            is_first_header = text == self._extract_text(block).strip()
            
            # Extract formatting from spans
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    size = span.get("size", 0)
                    if size > max_size:
                        max_size = size
                    
                    if span.get("flags", 0) & 16:  # Bold flag
                        is_bold = True
            
            # Determine header level
            body_size = layout_info.get('body_font_size', 11)
            if max_size <= body_size * 1.1:  # Slightly larger than body
                header_level = 3
            elif max_size <= body_size * 1.5:  # Moderately larger
                header_level = 2
            else:  # Significantly larger
                header_level = 1
            
            # Main title is likely the largest text on first page
            if is_first_header and max_size == max(layout_info.get('font_sizes', {0: 0}).keys(), default=0):
                header_level = 1
                
            # All uppercase text is often a header
            if text.isupper() and len(text) > 3:
                header_level = min(header_level, 2)  # Make sure it's at least level 2
            
            # Apply appropriate Word style based on header level
            if header_level == 1:
                paragraph.style = "Heading 1"
            elif header_level == 2:
                paragraph.style = "Heading 2"
            else:
                paragraph.style = "Heading 3"
            
            # Add the text with appropriate formatting
            run = paragraph.add_run(text)
            
            # Custom formatting beyond the style
            if header_level == 1:  # Main headers might need extra emphasis
                run.bold = True
                if max_size > 0:
                    run.font.size = Pt(max_size)  # Preserve original size for main headers
            
            # Underline section headers if they appear to be underlined in original
            if header_level == 2:
                # Try to detect if this header had a line under it
                try:
                    self._add_bottom_border(paragraph)
                except:
                    # If adding border fails, use style
                    pass
            
            # Set paragraph alignment (detect from original)
            alignment = self._detect_text_alignment(block)
            if alignment:
                paragraph.alignment = alignment
            
            # Add appropriate spacing after header
            if header_level == 1:
                paragraph.paragraph_format.space_after = Pt(12)
            elif header_level == 2:
                paragraph.paragraph_format.space_after = Pt(8)
            else:
                paragraph.paragraph_format.space_after = Pt(6)
        except Exception as e:
            logger.warning(f"Error formatting header: {str(e)}")
            # Fallback to simple text
            paragraph.text = self._extract_text(block).strip()
    
    def _add_bottom_border(self, paragraph):
        """Add bottom border to paragraph for header underlining"""
        try:
            # Access paragraph's XML element
            p = paragraph._p
            pPr = p.get_or_add_pPr()
            
            # Create border element
            pBdr = OxmlElement('w:pBdr')
            
            # Create bottom border element
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')  # Border width
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), '000000')  # Border color (black)
            
            # Add bottom border to border element
            pBdr.append(bottom)
            
            # Add border element to paragraph properties
            pPr.append(pBdr)
        except Exception as e:
            logger.debug(f"Error adding border: {str(e)}")
    
    def _process_content_block(self, cell, block, layout_info, current_section, is_list, is_resume=False):
        """Process a content block"""
        try:
            # Check if this is a skills/ratings block
            if layout_info['has_ratings'] and self._contains_ratings(block):
                self._process_ratings_block(cell, block)
                return
            
            # Create paragraph
            paragraph = cell.add_paragraph()
            
            # Check for bullet points
            text = self._extract_text(block)
            contains_bullets = any(bullet in text for bullet in layout_info['bullet_styles'])
            
            # If this looks like a list item but no bullet markers detected,
            # check for numerical markers (1., 2., etc.)
            if not contains_bullets and re.match(r'^\s*\d+\.\s', text):
                paragraph.style = "List Number"
                contains_bullets = True
            elif contains_bullets:
                paragraph.style = "List Bullet"
            
            # Check for a link line
            is_link = False
            if text.strip().startswith('http') or text.strip().startswith('www'):
                is_link = True
                # Add hyperlink style
                paragraph.style = "Hyperlink"
            
            # Special handling for resume items like email, phone, etc.
            if is_resume:
                # Try to detect contact info
                if re.match(r'^[\w\.-]+@[\w\.-]+\.\w+$', text.strip()):
                    # Email address
                    paragraph.style = "Hyperlink"
                    is_link = True
                elif re.match(r'^\+?[\d\s\(\)-]{7,}$', text.strip()):
                    # Phone number
                    paragraph.add_run(text)
                    paragraph.paragraph_format.space_after = Pt(2)
                    return
            
            # Process lines in the block
            for line_idx, line in enumerate(block.get("lines", [])):
                if line_idx > 0:
                    # Add line break between lines
                    paragraph.add_run().add_break()
                
                # Extract text for this line
                line_text = "".join(span.get("text", "") for span in line.get("spans", []))
                
                # Check if this is a bullet point
                is_bullet = re.match(r'^\s*[•⦁◦·○●◆⬤✦]\s', line_text) or re.match(r'^\s*\d+\.\s', line_text)
                
                if is_bullet:
                    # Handle as bullet point (strip the bullet character)
                    bullet_char = is_bullet.group(0)
                    clean_text = line_text.replace(bullet_char, "", 1).strip()
                    
                    # Use Word's built-in bullets
                    if line_idx == 0:  # Only set list style on first line
                        paragraph.style = "List Bullet" if not re.match(r'^\s*\d+\.\s', line_text) else "List Number"
                    
                    # Add the text without the bullet
                    run = paragraph.add_run(clean_text)
                    
                    # Preserve formatting from original
                    if line.get("spans"):
                        self._apply_span_formatting(run, line.get("spans")[0])
            else:
                    # Regular text - preserve spans and their formatting
                    spans = line.get("spans", [])
                    for span in spans:
                        span_text = span.get("text", "").strip()
                        if not span_text:
                            continue
                        
                        run = paragraph.add_run(span_text + " ")
                        self._apply_span_formatting(run, span)
                    
                    # If no spans were processed, add the line text directly
                    if not spans:
                        run = paragraph.add_run(line_text)
            
            # Set paragraph alignment
            alignment = self._detect_text_alignment(block)
            if alignment:
                paragraph.alignment = alignment
            
            # Set appropriate line and paragraph spacing
            if contains_bullets:
                # List items usually have less spacing
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.1
            else:
                # Regular paragraphs
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.1
                
            # Special handling for hyperlinks
            if is_link:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(5, 99, 193)  # Hyperlink blue
                    run.underline = True
                    
            # Special resume spacing
            if is_resume:
                paragraph.paragraph_format.space_before = Pt(1)
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.05
        except Exception as e:
            logger.warning(f"Error processing content block: {str(e)}")
            # Add a simple paragraph with the text to ensure content is not lost
            try:
                cell.add_paragraph(self._extract_text(block))
            except:
                pass
    
    def _contains_ratings(self, block):
        """Check if a block contains skill ratings (patterns of dots or circles)"""
        try:
            text = self._extract_text(block)
            # Look for patterns like ●●●○○ or ★★★☆☆
            return bool(re.search(r'[●○★☆■□]{3,}', text))
        except:
            return False
    
    def _process_ratings_block(self, cell, block):
        """Process a block containing skill ratings"""
        try:
            # Create paragraph for each line
            for line in block.get("lines", []):
                paragraph = cell.add_paragraph()
                
                # Get text for this line
                line_text = "".join(span.get("text", "") for span in line.get("spans", []))
                
                # Try to extract skill name and rating
                rating_match = re.search(r'(.*?)([●○★☆■□]{3,})', line_text)
                
                if rating_match:
                    # Split into skill name and rating
                    skill_name = rating_match.group(1).strip()
                    rating = rating_match.group(2)
                    
                    # Add skill name
                    name_run = paragraph.add_run(skill_name)
                    name_run.bold = True
                    
                    # Add appropriate spacing
                    space_length = max(1, min(20, 20 - len(skill_name)))
                    space_run = paragraph.add_run(" " * space_length)
                    
                    # Add rating symbols
                    rating_run = paragraph.add_run(rating)
                    
                    # Format rating appropriately
                    rating_run.font.name = "Segoe UI Symbol"
                    rating_run.font.size = Pt(10)
                else:
                    # Just add the text as-is
                    paragraph.add_run(line_text)
                
                # Add small spacing after rating line
                paragraph.paragraph_format.space_after = Pt(2)
        except Exception as e:
            logger.warning(f"Error processing ratings: {str(e)}")
            # Fallback to simple text
            try:
                cell.add_paragraph(self._extract_text(block))
            except:
                pass
    
    def _apply_span_formatting(self, run, span):
        """Apply text formatting from span to run"""
        try:
            # Set font name if available
            font_name = span.get("font", "Calibri")
            if font_name and len(font_name) > 0:
                run.font.name = font_name
            
            # Set font size
            size = span.get("size", 11)
            if size > 0:
                run.font.size = Pt(max(6, min(72, size)))  # Clamp between 6 and 72
            
            # Set font color if available
            if "color" in span:
                color = span["color"]
                # Convert from integer RGB values to RGBColor
                if isinstance(color, int):
                    # Extract RGB components (assuming color is stored as RGB integer)
                    r = (color >> 16) & 0xFF
                    g = (color >> 8) & 0xFF
                    b = color & 0xFF
                    run.font.color.rgb = RGBColor(r, g, b)
            
            # Apply text decorations
            flags = span.get("flags", 0)
            run.bold = bool(flags & 16)  # 2^4 = 16 (bold flag)
            run.italic = bool(flags & 2)  # 2^1 = 2 (italic flag)
            run.underline = bool(flags & 4)  # 2^2 = 4 (underline flag)
        except Exception as e:
            logger.debug(f"Error applying formatting: {str(e)}")
    
    def _extract_text(self, block):
        """Extract all text from a block"""
        try:
            text = ""
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text += span.get("text", "")
            return text
        except Exception as e:
            logger.debug(f"Error extracting text: {str(e)}")
            return ""
    
    def _detect_text_alignment(self, block):
        """Detect the alignment of text in a block"""
        try:
            # Get all spans in the block
            spans = []
            for line in block.get("lines", []):
                spans.extend(line.get("spans", []))
            
            if not spans:
                return None
            
            # Get page width and block position
            page_width = block.get("page", {}).get("width", 0)
            if not page_width:  # If page width not available in block
                page_width = 612  # Default A4 width in points
            
            if "bbox" not in block or len(block["bbox"]) < 4:
                return WD_PARAGRAPH_ALIGNMENT.LEFT
                
            block_left = block["bbox"][0]
            block_right = block["bbox"][2]
            block_width = block_right - block_left
            page_center = page_width / 2
            block_center = block_left + (block_width / 2)
            
            # Determine alignment by position
            left_margin = block_left
            right_margin = page_width - block_right
            
            # Center detection
            if abs(block_center - page_center) < 50:
                return WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Right alignment detection
            if left_margin > right_margin + 100:
                return WD_PARAGRAPH_ALIGNMENT.RIGHT
            
            # Default to left alignment
            return WD_PARAGRAPH_ALIGNMENT.LEFT
        except Exception as e:
            logger.debug(f"Error detecting alignment: {str(e)}")
            # Default to None if detection fails
            return WD_PARAGRAPH_ALIGNMENT.LEFT
    
    def _extract_and_add_images(self, doc, page):
        """Extract images from the PDF page and add them to the document"""
        try:
            # Get list of image blocks
            image_list = page.get_images(full=True)
            
            for img_index, img_info in enumerate(image_list):
                if not img_info or len(img_info) < 1:
                    continue
                    
                xref = img_info[0]  # Image reference number
                
                # Extract the image
                base_image = page.parent.extract_image(xref)
                image_bytes = base_image.get("image")
                
                # Get image rectangle (position and size)
                image_rect = None
                for block in page.get_text("dict").get("blocks", []):
                    if block.get("type") == 1 and block.get("xref") == xref:  # Image block
                        image_rect = block.get("bbox")
                        break
                
                if image_bytes:
                    # Create a temporary file for the image
                    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_file:
                        temp_file.write(image_bytes)
                        temp_path = temp_file.name
                    
                    # Add a paragraph for image positioning
                    p = doc.add_paragraph()
                    
                    # Set alignment based on image position
                    if image_rect:
                        x_center = (image_rect[0] + image_rect[2]) / 2
                        page_center = page.rect.width / 2
                        
                        if abs(x_center - page_center) < 50:  # Image is centered
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        elif x_center < page_center:  # Image is left-aligned
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        else:  # Image is right-aligned
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    
                    # Add the image
                    run = p.add_run()
                    run.add_picture(temp_path)
                    
                    # Clean up temporary file
                    try:
                        os.unlink(temp_path)
                    except:
                        pass
        
        except Exception as e:
            logger.warning(f"Error extracting images: {str(e)}")

    def _detect_alignment_zones(self, x_coordinates, page_width):
        """Detect alignment zones (left, center, right) on the page"""
        try:
            # Create a histogram of x-coordinates
            hist, bin_edges = np.histogram(x_coordinates, bins=min(50, len(x_coordinates)//5 + 5))
            
            # Identify peaks in the histogram (alignment zones)
            peaks = self._find_peaks(hist)
            
            # Convert peak indices to x-coordinates
            alignment_zones = [bin_edges[i] for i in peaks if i < len(bin_edges)]
            
            # Ensure we have at least common alignment zones
            if len(alignment_zones) < 2:
                # Add default left margin if missing
                left_margin = page_width * 0.1
                if not any(abs(zone - left_margin) < page_width * 0.05 for zone in alignment_zones):
                    alignment_zones.append(left_margin)
                
                # Add default right margin if missing
                right_margin = page_width * 0.9
                if not any(abs(zone - right_margin) < page_width * 0.05 for zone in alignment_zones):
                    alignment_zones.append(right_margin)
                    
                # Sort alignment zones
                alignment_zones.sort()
            
            return alignment_zones
        except Exception as e:
            logger.warning(f"Error detecting alignment zones: {str(e)}")
            # Return some default alignment zones
            return [page_width * 0.1, page_width * 0.5, page_width * 0.9]

    def _find_peaks(self, histogram):
        """Find peaks in a histogram"""
        try:
            # Minimum height for a peak
            min_height = max(histogram) * 0.2
            
            # Minimum distance between peaks (in bins)
            min_distance = max(1, len(histogram) // 10)
            
            # Find peaks
            peaks = []
            for i in range(1, len(histogram) - 1):
                # Check if this point is higher than its neighbors
                if histogram[i] > histogram[i-1] and histogram[i] > histogram[i+1]:
                    # Check minimum height
                    if histogram[i] >= min_height:
                        # Check minimum distance from existing peaks
                        if not any(abs(i - p) < min_distance for p in peaks):
                            peaks.append(i)
            
            return peaks
        except Exception as e:
            logger.debug(f"Error finding peaks: {str(e)}")
            return []
    
    def _detect_columns_advanced(self, x_coordinates, page_width):
        """Advanced column detection using density-based analysis"""
        try:
            if len(x_coordinates) < 5:
                return []
                
            # Create a density histogram (more bins for better resolution)
            num_bins = min(100, max(20, len(x_coordinates) // 3))
            hist, bin_edges = np.histogram(x_coordinates, bins=num_bins, range=(0, page_width))
            
            # Smooth the histogram to reduce noise
            smoothed_hist = np.convolve(hist, np.ones(3)/3, mode='same')
            
            # Find density valleys (spaces between columns)
            valleys = []
            for i in range(1, len(smoothed_hist) - 1):
                if (smoothed_hist[i] < smoothed_hist[i-1] and 
                    smoothed_hist[i] < smoothed_hist[i+1] and
                    smoothed_hist[i] < max(smoothed_hist) * 0.2):  # 20% of max height
                    valleys.append(i)
            
            # Convert valley indices to x-coordinates
            valley_x = [bin_edges[i] for i in valleys if i < len(bin_edges)]
            
            # Filter valleys that are too close together
            min_column_width = page_width * 0.1  # 10% of page width
            filtered_valleys = []
            for i, x in enumerate(valley_x):
                if i == 0 or x - valley_x[i-1] >= min_column_width:
                    filtered_valleys.append(x)
            
            # Create columns from valley boundaries
            columns = []
            if filtered_valleys:
                # Add left edge as first column boundary
                boundaries = [0] + filtered_valleys + [page_width]
                
                for i in range(len(boundaries) - 1):
                    left = boundaries[i]
                    right = boundaries[i+1]
                    width_ratio = (right - left) / page_width
                    
                    # Skip if column is too narrow
                    if width_ratio < 0.1:
                        continue
                        
                    columns.append({
                        'left': left,
                        'right': right,
                        'width_ratio': width_ratio
                    })
            
            return columns
        except Exception as e:
            logger.warning(f"Error in advanced column detection: {str(e)}")
            return []
    
    def _columns_from_alignment_zones(self, alignment_zones, page_width):
        """Create column definitions from alignment zones"""
        try:
            if not alignment_zones or len(alignment_zones) < 2:
                return []
                
            # Sort alignment zones
            zones = sorted(alignment_zones)
            
            # Create columns
            columns = []
            
            # First column from page edge to first zone
            if zones[0] > page_width * 0.05:  # Only if first zone isn't at page edge
                width_ratio = zones[0] / page_width
                columns.append({
                    'left': 0,
                    'right': zones[0],
                    'width_ratio': width_ratio
                })
            
            # Middle columns between zones
            for i in range(len(zones) - 1):
                left = zones[i]
                right = zones[i+1]
                width_ratio = (right - left) / page_width
                
                # Only add if column has reasonable width
                if width_ratio >= 0.1:
                    columns.append({
                        'left': left,
                        'right': right,
                        'width_ratio': width_ratio
                    })
            
            # Last column from last zone to page edge
            if page_width - zones[-1] > page_width * 0.05:  # Only if there's space after last zone
                width_ratio = (page_width - zones[-1]) / page_width
                columns.append({
                    'left': zones[-1],
                    'right': page_width,
                    'width_ratio': width_ratio
                })
            
            return columns
        except Exception as e:
            logger.warning(f"Error creating columns from alignment zones: {str(e)}")
            return []

    def _analyze_page_layout(self, page):
        """Advanced page layout analysis with improved structure detection"""
        try:
            # Extract text blocks and prepare for analysis
            blocks = page.get_text("dict").get("blocks", [])
            text_blocks = [b for b in blocks if b.get("type") == 0]
            
            # Initialize layout info with more detailed structure
            layout_info = {
                'type': 'single_column',  # Default to single column
                'columns': [],
                'headers': [],
                'has_ratings': False,
                'bullet_styles': set(),
                'page_width': page.rect.width,
                'page_height': page.rect.height,
                'font_sizes': {},
                'all_x_coordinates': [],
                'alignment_zones': [],
                'text_blocks_count': len(text_blocks),
                'is_resume': False,
                'text_density_map': {},  # NEW: Track text density across the page
                'margins': {            # NEW: Track document margins
                    'left': page.rect.width,
                    'right': 0,
                    'top': page.rect.height,
                    'bottom': 0
                },
                'line_heights': [],     # NEW: Track line heights for better spacing
                'indentation_levels': [],# NEW: Track paragraph indentations
                'vertical_spacing': []   # NEW: Track vertical spacing between elements
            }
            
            # More precise text block analysis
            x_coordinates = []
            y_coordinates = []
            
            # Create a density map of the page (divide into 100x100 grid)
            density_map = {}
            grid_size = 100
            for i in range(grid_size):
                for j in range(grid_size):
                    density_map[f"{i}_{j}"] = 0
                    
            page_width = page.rect.width
            page_height = page.rect.height
            
            # Extract comprehensive layout information
            for block in text_blocks:
                if "bbox" in block and len(block["bbox"]) >= 4:
                    # Extract precise coordinates
                    x_min = block["bbox"][0]
                    y_min = block["bbox"][1]
                    x_max = block["bbox"][2]
                    y_max = block["bbox"][3]
                    
                    # Track boundaries for accurate margin detection
                    layout_info['margins']['left'] = min(layout_info['margins']['left'], x_min)
                    layout_info['margins']['right'] = max(layout_info['margins']['right'], x_max)
                    layout_info['margins']['top'] = min(layout_info['margins']['top'], y_min)
                    layout_info['margins']['bottom'] = max(layout_info['margins']['bottom'], y_max)
                    
                    # Collect coordinate data
                    x_coordinates.append(x_min)
                    x_coordinates.append(x_max)
                    layout_info['all_x_coordinates'].append(x_min)
                    layout_info['all_x_coordinates'].append(x_max)
                    
                    y_coordinates.append(y_min)
                    y_coordinates.append(y_max)
                    
                    # Update density map
                    width = x_max - x_min
                    height = y_max - y_min
                    
                    # Convert to grid coordinates
                    grid_x_min = int((x_min / page_width) * grid_size)
                    grid_x_max = int((x_max / page_width) * grid_size)
                    grid_y_min = int((y_min / page_height) * grid_size)
                    grid_y_max = int((y_max / page_height) * grid_size)
                    
                    # Update density in all overlapping grid cells
                    for i in range(max(0, grid_x_min), min(grid_size, grid_x_max + 1)):
                        for j in range(max(0, grid_y_min), min(grid_size, grid_y_max + 1)):
                            density_map[f"{i}_{j}"] += 1
                
                # Track line heights for better spacing detection
                for line_idx, line in enumerate(block.get("lines", [])):
                    if "bbox" in line and len(line["bbox"]) >= 4:
                        line_height = line["bbox"][3] - line["bbox"][1]
                        layout_info['line_heights'].append(line_height)
                        
                        # Track indentation level
                        if line_idx == 0 and line["bbox"][0] > layout_info['margins']['left'] + 5:
                            # This might be an indented paragraph
                            indentation = line["bbox"][0] - layout_info['margins']['left']
                            layout_info['indentation_levels'].append(indentation)
                
                # Check for vertical spacing between blocks
                if len(y_coordinates) >= 4:  # Need at least two blocks to compute spacing
                    prev_block_bottom = y_coordinates[-3]  # Second-to-last bottom
                    curr_block_top = y_coordinates[-2]     # Last top
                    spacing = curr_block_top - prev_block_bottom
                    if spacing > 0:
                        layout_info['vertical_spacing'].append(spacing)
                
                # Collect formatting information for design analysis
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        font_size = span.get("size", 0)
                        font_name = span.get("font", "")
                        text_color = span.get("color", 0)
                        
                        # Track font sizes with frequency
                        layout_info['font_sizes'][font_size] = layout_info['font_sizes'].get(font_size, 0) + 1
                    
                    # Check for bullet points with more patterns
                    line_text = "".join(span.get("text", "") for span in line.get("spans", []))
                    bullet_patterns = [
                        r'^\s*[•⦁◦·○●◆⬤✦]\s',  # Common bullet symbols
                        r'^\s*[-–—]\s',         # Dashes as bullets
                        r'^\s*\d+\.\s',         # Numbered lists
                        r'^\s*[a-zA-Z]\)\s',    # Letter lists like a) b) c)
                        r'^\s*\(\d+\)\s'        # Number in parentheses like (1) (2)
                    ]
                    
                    for pattern in bullet_patterns:
                        if re.match(pattern, line_text):
                            # Extract the actual bullet character used
                            bullet_match = re.match(pattern, line_text)
                            if bullet_match and bullet_match.group(0):
                                bullet_char = bullet_match.group(0).strip()
                                if bullet_char:
                                    layout_info['bullet_styles'].add(bullet_char[0])
            
            # Store density map
            layout_info['text_density_map'] = density_map
            
            # Enhance document type detection
            text_content = page.get_text().upper()
            resume_keywords = [
                'RESUME', 'CV', 'CURRICULUM VITAE', 'PROFILE', 'EXPERIENCE', 'EDUCATION', 'SKILLS',
                'WORK HISTORY', 'EMPLOYMENT', 'KONTAKT', 'PROFIL', 'UTBILDNING', 'ARBETSLIVSERFARENHET',
                'SUMMARY', 'OBJECTIVE', 'QUALIFICATIONS', 'CAREER HIGHLIGHTS', 'CERTIFICATIONS'
            ]
            
            # Check for typical resume patterns
            has_contact_info = bool(re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text_content))  # Email
            has_phone = bool(re.search(r'(?:\+|00)?[0-9()\s-]{7,}', text_content))  # Phone number
            
            # Count resume keywords found
            resume_keyword_count = sum(1 for keyword in resume_keywords if keyword in text_content)
            
            # If we found multiple resume keywords OR contact info with at least one keyword
            if resume_keyword_count >= 2 or (resume_keyword_count >= 1 and (has_contact_info or has_phone)):
                layout_info['is_resume'] = True
                
            # IMPROVED: Advanced column detection using multiple methods
            
            # 1. First try density map analysis for column detection
            columns_from_density = self._detect_columns_from_density(density_map, grid_size, page.rect.width)
            
            # 2. Then try improved histogram method
            columns_from_histogram = []
            if len(x_coordinates) > 5:  # Need enough blocks for reliable detection
                try:
                    columns_from_histogram = self._detect_columns_advanced(x_coordinates, page.rect.width)
                except Exception as e:
                    logger.warning(f"Enhanced column detection failed: {str(e)}")
            
            # 3. Finally try alignment zone analysis
            columns_from_alignment = []
            if layout_info['all_x_coordinates']:
                try:
                    alignment_zones = self._detect_alignment_zones(np.array(layout_info['all_x_coordinates']), page.rect.width)
                    layout_info['alignment_zones'] = alignment_zones
                    columns_from_alignment = self._columns_from_alignment_zones(alignment_zones, page.rect.width)
                except Exception as e:
                    logger.warning(f"Alignment zones detection failed: {str(e)}")
            
            # Choose the best column detection method based on results
            all_column_results = [
                (columns_from_density, "density"),
                (columns_from_histogram, "histogram"),
                (columns_from_alignment, "alignment")
            ]
            
            # Filter out empty results
            valid_column_results = [(cols, method) for cols, method in all_column_results if cols and len(cols) > 0]
            
            if valid_column_results:
                # Prefer methods with more columns detected, as that likely indicates a more precise analysis
                valid_column_results.sort(key=lambda x: len(x[0]), reverse=True)
                best_columns, best_method = valid_column_results[0]
                
                if len(best_columns) > 1:
                    layout_info['type'] = 'multi_column'
                    layout_info['columns'] = best_columns
                    layout_info['column_detection_method'] = best_method
                    logger.debug(f"Detected {len(best_columns)} columns using {best_method} method")
                else:
                    layout_info['type'] = 'single_column'
                    layout_info['columns'] = []
                    logger.debug("Using single column layout")
            else:
                # Fall back to single column
                layout_info['type'] = 'single_column'
                layout_info['columns'] = []
                logger.debug("Falling back to single column layout (no columns detected)")
            
            # Find most common paragraph spacing for consistent layout
            if layout_info['vertical_spacing']:
                # Use statistical mode to find most common spacing
                from statistics import mode, mean
                try:
                    layout_info['common_paragraph_spacing'] = mode(layout_info['vertical_spacing'])
                except StatisticsError:
                    # If no unique mode, use mean
                    layout_info['common_paragraph_spacing'] = mean(layout_info['vertical_spacing'])
            
            # Identify headers (larger or bold text) with improved detection
            layout_info['headers'] = self._detect_headers(text_blocks, layout_info['font_sizes'])
            
            # Check for rating indicators (dots or stars)
            layout_info['has_ratings'] = self._detect_ratings(text_blocks)
            
            return layout_info
            
        except Exception as e:
            logger.warning(f"Error analyzing page layout: {str(e)}")
            # Return default layout info if analysis fails
            return {
                'type': 'single_column',
                'columns': [],
                'headers': [],
                'has_ratings': False,
                'bullet_styles': set(),
                'page_width': page.rect.width if page else 612,  # Default A4 width
                'page_height': page.rect.height if page else 792,  # Default A4 height
                'font_sizes': {},
                'all_x_coordinates': [],
                'text_blocks_count': 0
            }
            
    def _detect_columns_from_density(self, density_map, grid_size, page_width):
        """Detect columns using the text density map"""
        try:
            # Create a vertical density profile
            vertical_profile = [0] * grid_size
            
            for i in range(grid_size):
                for j in range(grid_size):
                    vertical_profile[i] += density_map.get(f"{i}_{j}", 0)
            
            # Find valleys in the density profile (these are gaps between columns)
            valleys = []
            for i in range(1, grid_size - 1):
                # Check if this is a local minimum
                if (vertical_profile[i] < vertical_profile[i-1] and 
                    vertical_profile[i] < vertical_profile[i+1] and
                    vertical_profile[i] < max(vertical_profile) * 0.2):  # Significant gap
                    valleys.append(i)
            
            # Convert to page coordinates
            valley_positions = [(v * page_width) / grid_size for v in valleys]
            
            # Need at least one valley to create columns
            if not valley_positions:
                return []
                
            # Create columns from valley boundaries
            columns = []
            
            # Add left edge as first boundary
            boundaries = [0] + valley_positions + [page_width]
            
            for i in range(len(boundaries) - 1):
                left = boundaries[i]
                right = boundaries[i+1]
                width_ratio = (right - left) / page_width
                
                # Skip if column is too narrow
                if width_ratio < 0.1:  # Less than 10% of page width
                    continue
                    
                columns.append({
                    'left': left,
                    'right': right,
                    'width_ratio': width_ratio
                })
            
            return columns
            
        except Exception as e:
            logger.warning(f"Error in density-based column detection: {str(e)}")
            return []
    
    def _detect_headers(self, text_blocks, font_sizes):
        """Improved header detection with multiple signals"""
        try:
            # Find the most common (body) font size
            if not font_sizes:
                return []
                
            body_font_size = max(font_sizes.items(), key=lambda x: x[1])[0]
            
            # Headers are typically larger or bold
            headers = []
            
            for block in text_blocks:
                # Extract block text once for efficiency
                block_text = self._extract_text(block).strip()
                if not block_text:
                    continue
                    
                # Skip very long text blocks (likely not headers)
                if len(block_text) > 100:
                    continue
                
                is_header = False
                max_size = 0
                is_bold = False
                all_caps = block_text.isupper() and len(block_text) > 2
                
                # Check for header-like formatting
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        size = span.get("size", 0)
                        if size > max_size:
                            max_size = size
                        
                        # Check for bold flag (16)
                        if span.get("flags", 0) & 16:
                            is_bold = True
                
                # Combine signals to detect headers
                if (max_size > body_font_size * 1.1) or (is_bold and all_caps):
                    is_header = True
                    
                # Additional checks for section headers
                if block_text.endswith(':') and len(block_text) < 30:
                    is_header = True
                    
                # Resume section headers often match these patterns
                resume_section_patterns = [
                    r'^(EDUCATION|EXPERIENCE|SKILLS|WORK HISTORY|EMPLOYMENT|PROFILE|SUMMARY|OBJECTIVE|QUALIFICATIONS)$',
                    r'^(UTBILDNING|ARBETSLIVSERFARENHET|FÄRDIGHETER|PROFIL)$'  # Swedish
                ]
                
                for pattern in resume_section_patterns:
                    if re.match(pattern, block_text, re.IGNORECASE):
                        is_header = True
                        break
                        
                if is_header:
                    header_level = 1  # Default level
                    
                    # Determine header level based on size and formatting
                    if max_size <= body_font_size * 1.1:  # Slightly larger than body
                        header_level = 3
                    elif max_size <= body_font_size * 1.5:  # Moderately larger
                        header_level = 2
                    else:  # Significantly larger
                        header_level = 1
                        
                    headers.append({
                        'text': block_text,
                        'level': header_level,
                        'position': block.get("bbox", [0, 0, 0, 0]),
                        'is_bold': is_bold,
                        'is_allcaps': all_caps,
                        'size': max_size
                    })
            
            return headers
            
        except Exception as e:
            logger.warning(f"Error detecting headers: {str(e)}")
            return []
            
    def _detect_ratings(self, text_blocks):
        """Detect if document contains rating indicators (skill bars, stars, etc.)"""
        try:
            for block in text_blocks:
                text = self._extract_text(block)
                
                # Check for rating patterns
                rating_patterns = [
                    r'[●○★☆■□]{3,}',    # Filled/empty circles or stars
                    r'[\|│]{3,}',       # Vertical bars (skill meter)
                    r'[▮▯]{2,}',        # Block indicators
                    r'[0-9]+\s*\/\s*[0-9]+',  # Numeric ratings (e.g., 8/10)
                    r'[0-9]+\s*out of\s*[0-9]+'  # Text ratings (e.g., 4 out of 5)
                ]
                
                for pattern in rating_patterns:
                    if re.search(pattern, text):
                        return True
            
            return False
            
        except Exception as e:
            logger.debug(f"Error detecting ratings: {str(e)}")
            return False

    def _combine_layout_info(self, layout_info_all_pages):
        """Combine layout information from multiple pages to determine global document style"""
        try:
            if not layout_info_all_pages:
                return {}
                
            # Initialize global layout with default values
            global_layout = {
                'layout_type': 'single_column',
                'consistent_layout_type': False,
                'header_font_sizes': [],
                'body_font_size': 11,
                'base_font': 'Calibri',
                'bullet_styles': set(),
                'has_ratings': False,
                'is_resume': False
            }
            
            # Count layout types
            layout_types = {}
            for info in layout_info_all_pages:
                layout_type = info.get('type', 'single_column')
                layout_types[layout_type] = layout_types.get(layout_type, 0) + 1
            
            # Determine dominant layout type
            if layout_types:
                dominant_type, count = max(layout_types.items(), key=lambda x: x[1])
                global_layout['layout_type'] = dominant_type
                
                # Check if layout type is consistent across pages
                if count == len(layout_info_all_pages):
                    global_layout['consistent_layout_type'] = True
            
            # Combine font information
            all_font_sizes = {}
            for info in layout_info_all_pages:
                for size, count in info.get('font_sizes', {}).items():
                    all_font_sizes[size] = all_font_sizes.get(size, 0) + count
            
            # Determine body font size (most common size)
            if all_font_sizes:
                global_layout['body_font_size'] = max(all_font_sizes.items(), key=lambda x: x[1])[0]
                
                # Extract header sizes (larger than body text)
                header_sizes = [size for size in all_font_sizes.keys() 
                               if size > global_layout['body_font_size'] * 1.1]
                header_sizes.sort(reverse=True)  # Sort from largest to smallest
                global_layout['header_font_sizes'] = header_sizes[:3]  # Top 3 header sizes
            
            # Combine bullet styles
            for info in layout_info_all_pages:
                global_layout['bullet_styles'].update(info.get('bullet_styles', set()))
            
            # Check if any page has ratings
            global_layout['has_ratings'] = any(info.get('has_ratings', False) for info in layout_info_all_pages)
            
            # Check if document is likely a resume
            global_layout['is_resume'] = any(info.get('is_resume', False) for info in layout_info_all_pages)
            
            # Find most common base font
            font_counts = {}
            for info in layout_info_all_pages:
                for line in [line for block in info.get('blocks', []) 
                            if isinstance(block, dict) and 'lines' in block
                            for line in block['lines']]:
                    for span in line.get('spans', []):
                        font = span.get('font', '')
                        if font:
                            font_counts[font] = font_counts.get(font, 0) + 1
            
            if font_counts:
                global_layout['base_font'] = max(font_counts.items(), key=lambda x: x[1])[0]
            
            return global_layout
            
        except Exception as e:
            logger.warning(f"Error combining layout info: {str(e)}")
            return {
                'layout_type': 'single_column',
                'consistent_layout_type': False,
                'header_font_sizes': [16, 14, 12],
                'body_font_size': 11,
                'base_font': 'Calibri'
            }
            
    def _set_page_properties(self, section, page):
        """Set page size and orientation in document section"""
        try:
            # Get page dimensions
            page_width = page.rect.width
            page_height = page.rect.height
            
            # Convert PDF points to inches (1 point = 1/72 inch)
            width_inches = page_width / 72
            height_inches = page_height / 72
            
            # Set page dimensions
            section.page_width = Inches(width_inches)
            section.page_height = Inches(height_inches)
            
            # Set orientation
            if page_width > page_height:
                section.orientation = WD_ORIENTATION.LANDSCAPE
            else:
                section.orientation = WD_ORIENTATION.PORTRAIT
                
            # Set margins
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            
        except Exception as e:
            logger.warning(f"Error setting page properties: {str(e)}")
            # Use default properties if error occurs

# Keep SmartDocumentConverter for backward compatibility
class SmartDocumentConverter:
    def __init__(self, analysis=None):
        self.analysis = analysis
        self.reset_state()
        logger.debug("Using DocumentConverter for conversion")
        self.basic_converter = DocumentConverter()
        
    def reset_state(self):
        """Reset converter state for new conversion"""
        self.section_styles = {}
        self.current_structure = None
        self.temp_files = []
        
    def convert_document(self, input_path, output_path):
        """Fallback to basic converter"""
        try:
            # Use the basic converter instead
            return self.basic_converter.convert(input_path, output_path)
        except Exception as e:
            logger.error(f"Conversion error: {str(e)}")
            raise