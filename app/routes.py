from flask import Blueprint, render_template, request, send_file, current_app, redirect, url_for, flash, session, jsonify
from flask_login import login_required, current_user
import os
from werkzeug.utils import secure_filename
from deep_translator import GoogleTranslator
import docx
from pdf2docx import Converter as PDFConverter
import tempfile
from pathlib import Path
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import simpleSplit
from functools import wraps
from pdf2docx import Converter
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pdf2docx import parse
from PyPDF2 import PdfReader
from PIL import Image
import fitz  # PyMuPDF
import markdown
import subprocess
from docx2pdf import convert
import io

main = Blueprint('main', __name__)

def premium_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('is_premium'):
            flash('This feature requires a premium subscription', 'premium')
            return redirect(url_for('main.pricing'))
        return f(*args, **kwargs)
    return decorated_function

@main.route('/')
def index():
    """Route for the main PDF converter page"""
    try:
        translations = {
            'title': 'PDF to Word Converter',
            'security_title': 'Security Information',
            'security_info': [
                'Your files are encrypted during transfer',
                'Files are automatically deleted after conversion',
                'No data is stored on our servers'
            ]
        }
        
        return render_template('index.html',
            lang='en',
            translations=translations
        )
    except Exception as e:
        current_app.logger.error(f"Error rendering index: {str(e)}")
        return f"An error occurred: {str(e)}", 500

@main.route('/translate')
def translate_page():
    """Route for the document translator page"""
    try:
        languages = {
            'en': 'English',
            'es': 'Spanish',
            'fr': 'French',
            'de': 'German',
            'it': 'Italian',
            'pt': 'Portuguese',
            'ru': 'Russian',
            'ko': 'Korean',
            'ar': 'Arabic',
            'sv': 'Swedish'
        }
        
        return render_template('translate.html',
            languages=languages,
            translations={
                'title': 'Document Translator',
                'description': 'Translate your documents between languages',
                'upload_text': 'Drag & drop your document here or click to browse',
                'supported_formats': 'Supported formats: PDF, DOCX, TXT'
            }
        )
    except Exception as e:
        current_app.logger.error(f"Error rendering translate page: {str(e)}")
        return f"An error occurred: {str(e)}", 500

@main.route('/translate/convert', methods=['POST'])
def translate_document():
    try:
        if 'file' not in request.files:
            return 'No file uploaded', 400
        
        file = request.files['file']
        source_lang = request.form.get('source_lang', 'auto')
        target_lang = request.form.get('target_lang', 'en')
        
        if file.filename == '':
            return 'No file selected', 400
            
        # Create temp directory for processing
        with tempfile.TemporaryDirectory() as temp_dir:
            # Save the uploaded file
            filename = secure_filename(file.filename)
            input_path = os.path.join(temp_dir, filename)
            file.save(input_path)
            
            # Extract text based on file type
            text_content = ''
            file_extension = Path(filename).suffix.lower()
            
            if file_extension == '.txt':
                with open(input_path, 'r', encoding='utf-8') as f:
                    text_content = f.read()
                    
            elif file_extension == '.docx':
                doc = docx.Document(input_path)
                text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                
            elif file_extension == '.pdf':
                # Convert PDF to DOCX first
                temp_docx = os.path.join(temp_dir, 'temp.docx')
                cv = PDFConverter(input_path)
                cv.convert(temp_docx)
                cv.close()
                
                # Extract text from DOCX
                doc = docx.Document(temp_docx)
                text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            # Translate the content
            try:
                translator = GoogleTranslator(source=source_lang, target=target_lang)
                translated_text = translator.translate(text_content)
            except Exception as e:
                current_app.logger.error(f"Translation error: {str(e)}")
                return f"Translation failed: {str(e)}", 500
            
            # Create output file based on format
            if file_extension == '.pdf':
                final_output = os.path.join(temp_dir, f'translated_{filename}')
                try:
                    # Create PDF with ReportLab
                    c = canvas.Canvas(final_output, pagesize=letter)
                    width, height = letter
                    
                    # Use built-in Helvetica font
                    c.setFont('Helvetica', 12)
                    
                    # Add title
                    c.setFont('Helvetica-Bold', 16)
                    title = f"Translated Document ({source_lang} → {target_lang})"
                    c.drawString(50, height - 40, title)
                    
                    # Reset font for body text
                    c.setFont('Helvetica', 12)
                    
                    # Split text into lines and pages
                    y = height - 80  # Start below title
                    for paragraph in translated_text.split('\n'):
                        if paragraph.strip():
                            # Wrap long lines
                            lines = simpleSplit(paragraph, 'Helvetica', 12, width - 100)
                            for line in lines:
                                if y < 50:  # New page if near bottom
                                    c.showPage()
                                    c.setFont('Helvetica', 12)
                                    y = height - 50
                                try:
                                    c.drawString(50, y, line)
                                except:
                                    # If there's an encoding issue, try to clean the text
                                    cleaned_line = ''.join(char for char in line if ord(char) < 128)
                                    c.drawString(50, y, cleaned_line)
                                y -= 20  # Line spacing
                            y -= 10  # Paragraph spacing
                    
                    # Add page numbers
                    page_num = c.getPageNumber()
                    c.drawString(width/2, 30, f"Page {page_num}")
                    
                    c.save()
                except Exception as e:
                    current_app.logger.error(f"PDF creation error: {str(e)}")
                    return f"Error creating PDF: {str(e)}", 500
                    
            elif file_extension == '.docx':
                final_output = os.path.join(temp_dir, f'translated_{filename}')
                doc = docx.Document()
                for paragraph in translated_text.split('\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph)
                doc.save(final_output)
                
            else:  # .txt
                final_output = os.path.join(temp_dir, f'translated_{filename}')
                with open(final_output, 'w', encoding='utf-8') as f:
                    f.write(translated_text)
            
            # Send the file
            try:
                return send_file(
                    final_output,
                    as_attachment=True,
                    download_name=f'translated_{filename}',
                    mimetype='application/pdf' if file_extension == '.pdf' else None
                )
            except Exception as e:
                current_app.logger.error(f"Error sending file: {str(e)}")
                return f"Error sending file: {str(e)}", 500
                
    except Exception as e:
        current_app.logger.error(f"Translation error: {str(e)}")
        return f"Translation failed: {str(e)}", 500

ALLOWED_EXTENSIONS = {'pdf'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def convert_document(file_path, output_format):
    try:
        output_path = os.path.splitext(file_path)[0] + f'.{output_format}'
        
        if output_format == 'docx':
            # Existing DOCX conversion
            pdf = PdfReader(file_path)
            total_pages = len(pdf.pages)
            parse(file_path, output_path, start=0, end=total_pages)
            
            # Set document properties
            doc = Document(output_path)
            doc.core_properties.author = "PDF Converter"
            doc.core_properties.title = os.path.basename(file_path)
            doc.save(output_path)
            
        elif output_format == 'txt':
            # Convert PDF to text
            pdf = PdfReader(file_path)
            with open(output_path, 'w', encoding='utf-8') as txt_file:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        txt_file.write(text + '\n\n')

        elif output_format == 'rtf':
            # Convert to DOCX first, then to RTF
            temp_docx = output_path + '.temp.docx'
            parse(file_path, temp_docx)
            
            # Use soffice (LibreOffice) to convert DOCX to RTF
            subprocess.run(['soffice', '--headless', '--convert-to', 'rtf', '--outdir', 
                          os.path.dirname(output_path), temp_docx])
            
            # Cleanup temp file
            if os.path.exists(temp_docx):
                os.remove(temp_docx)

        elif output_format == 'xlsx':
            # Convert PDF to XLSX using Tabula
            from tabula import read_pdf
            df = read_pdf(file_path, pages='all')
            if isinstance(df, list):
                import pandas as pd
                writer = pd.ExcelWriter(output_path, engine='openpyxl')
                for i, table in enumerate(df):
                    table.to_excel(writer, sheet_name=f'Sheet{i+1}', index=False)
                writer.close()

        elif output_format == 'pptx':
            from pdf2pptx import convert_pdf_to_pptx
            convert_pdf_to_pptx(file_path, output_path)

        elif output_format == 'epub':
            # Convert to HTML first
            temp_html = output_path + '.temp.html'
            pdf = fitz.open(file_path)
            html_content = ""
            for page in pdf:
                html_content += page.get_text("html")
            with open(temp_html, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # Convert HTML to EPUB using Pandoc
            subprocess.run(['pandoc', '-f', 'html', '-t', 'epub', '-o', output_path, temp_html])
            
            # Cleanup
            if os.path.exists(temp_html):
                os.remove(temp_html)

        elif output_format == 'html':
            pdf = fitz.open(file_path)
            html_content = ""
            for page in pdf:
                html_content += page.get_text("html")
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)

        elif output_format == 'md':
            # Convert to HTML first, then to Markdown
            pdf = fitz.open(file_path)
            html_content = ""
            for page in pdf:
                html_content += page.get_text("html")
            
            # Convert HTML to Markdown
            from html2text import html2text
            md_content = html2text(html_content)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(md_content)

        elif output_format == 'odt':
            # Convert to DOCX first, then to ODT
            temp_docx = output_path + '.temp.docx'
            parse(file_path, temp_docx)
            
            # Use LibreOffice to convert to ODT
            subprocess.run(['soffice', '--headless', '--convert-to', 'odt', '--outdir', 
                          os.path.dirname(output_path), temp_docx])
            
            # Cleanup
            if os.path.exists(temp_docx):
                os.remove(temp_docx)

        elif output_format == 'jpeg':
            pdf = fitz.open(file_path)
            for i, page in enumerate(pdf):
                pix = page.get_pixmap()
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                if i == 0:  # Save first page only
                    img.save(output_path, 'JPEG', quality=95)
                    break

        else:
            raise ValueError(f"Conversion to {output_format} not yet implemented")
            
        return output_path
            
    except Exception as e:
        current_app.logger.error(f"Conversion error: {str(e)}")
        raise Exception(f"Conversion failed: {str(e)}")

@main.route('/convert', methods=['POST'])
def convert_file():
    try:
        if 'file' not in request.files:
            current_app.logger.error("No file in request")
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        output_format = request.form.get('format', 'docx')
        
        current_app.logger.info(f"Converting file: {file.filename} to {output_format}")
        
        if file and allowed_file(file.filename):
            original_name = os.path.splitext(file.filename)[0]
            uploads_dir = os.path.join(current_app.root_path, 'uploads')
            
            if not os.path.exists(uploads_dir):
                os.makedirs(uploads_dir)
            
            filename = secure_filename(file.filename)
            file_path = os.path.join(uploads_dir, filename)
            file.save(file_path)
            
            current_app.logger.info(f"File saved to: {file_path}")
            
            try:
                converted_file_path = convert_document(file_path, output_format)
                current_app.logger.info(f"File converted successfully: {converted_file_path}")
                
                return send_file(
                    converted_file_path,
                    as_attachment=True,
                    download_name=f'{original_name}.{output_format}'
                )
                
            except Exception as e:
                current_app.logger.error(f"Conversion error: {str(e)}")
                return jsonify({'error': str(e)}), 500
                
            finally:
                # Cleanup files
                if os.path.exists(file_path):
                    os.remove(file_path)
                if 'converted_file_path' in locals() and os.path.exists(converted_file_path):
                    os.remove(converted_file_path)
        
        current_app.logger.error(f"Invalid file: {file.filename}")
        return jsonify({'error': 'Invalid file'}), 400
        
    except Exception as e:
        current_app.logger.error(f"Unexpected error: {str(e)}")
        return jsonify({'error': 'An unexpected error occurred'}), 500

@main.route('/features')
def features():
    """Route for the features page"""
    try:
        return render_template('features.html', 
                             title='Features',
                             current_user=current_user)
    except Exception as e:
        current_app.logger.error(f"Error rendering features page: {str(e)}")
        return f"An error occurred: {str(e)}", 500

@main.route('/pricing')
def pricing():
    return render_template('pricing.html', title='Pricing')

@main.route('/signin')
def signin():
    return render_template('signin.html', title='Sign In')

@main.route('/signup')
def signup():
    return render_template('signup.html', title='Sign Up')

@main.route('/privacy')
def privacy():
    return render_template('privacy.html', title='Privacy Notice')

@main.route('/cookies')
def cookies():
    return render_template('cookies.html', title='Cookie and Advertising Notice')

@main.route('/about')
def about():
    return render_template('about.html', title='About Us')

@main.context_processor
def inject_now():
    return {'now': datetime.utcnow()}
