from flask import Blueprint, render_template, request, send_file, current_app, redirect, url_for, flash, session, jsonify
from flask_login import login_required, current_user
import os
from werkzeug.utils import secure_filename
from deep_translator import GoogleTranslator
import docx
from pdf2docx import Converter as PDFConverter
import tempfile
from pathlib import Path
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.utils import simpleSplit
from functools import wraps
from datetime import datetime
from . import main
import subprocess

# Update ALLOWED_EXTENSIONS to include all supported formats
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'rtf', 'xlsx', 'pptx', 'epub', 'html', 'md', 'odt', 'jpeg'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@main.route('/')
def index():
    return render_template('index.html')

@main.route('/features')
def features():
    return render_template('features.html')

@main.route('/pricing')
def pricing():
    return render_template('pricing.html')

@main.route('/signin')
def signin():
    return render_template('signin.html')

@main.route('/signup')
def signup():
    return render_template('signup.html')

@main.route('/about')
def about():
    return render_template('about.html')

@main.route('/privacy')
def privacy():
    return render_template('privacy.html')

@main.route('/terms')
def terms():
    return render_template('terms.html')

@main.route('/contact')
def contact():
    return render_template('contact.html')

@main.route('/cookies')
def cookies():
    return render_template('cookies.html')

@main.route('/translate')
def translate_page():
    try:
        languages = {
            'en': 'English',
            'es': 'Spanish',
            'fr': 'French',
            'de': 'German',
            'it': 'Italian',
            'pt': 'Portuguese',
            'ru': 'Russian',
            'ko': 'Korean',
            'ar': 'Arabic',
            'sv': 'Swedish'
        }
        
        return render_template('translate.html',
            languages=languages,
            translations={
                'title': 'Document Translator',
                'description': 'Translate your documents between languages',
                'upload_text': 'Drag & drop your document here or click to browse',
                'supported_formats': 'Supported formats: PDF, DOCX, TXT'
            }
        )
    except Exception as e:
        current_app.logger.error(f"Error rendering translate page: {str(e)}")
        return f"An error occurred: {str(e)}", 500

@main.route('/translate/convert', methods=['POST'])
def translate_document():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
            
        if not allowed_file(file.filename):
            return jsonify({'error': 'Invalid file type'}), 400

        source_lang = request.form.get('source_lang', 'auto')
        target_lang = request.form.get('target_lang', 'en')
        
        # Create temp directory for processing
        with tempfile.TemporaryDirectory() as temp_dir:
            filename = secure_filename(file.filename)
            input_path = os.path.join(temp_dir, filename)
            file.save(input_path)
            
            # Extract text based on file type
            text_content = ''
            file_extension = Path(filename).suffix.lower()
            
            if file_extension == '.txt':
                with open(input_path, 'r', encoding='utf-8') as f:
                    text_content = f.read()
            elif file_extension == '.docx':
                doc = docx.Document(input_path)
                text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            elif file_extension == '.pdf':
                temp_docx = os.path.join(temp_dir, 'temp.docx')
                cv = PDFConverter(input_path)
                cv.convert(temp_docx)
                cv.close()
                doc = docx.Document(temp_docx)
                text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            elif file_extension in ['.rtf', '.odt']:
                # Convert to DOCX first using LibreOffice
                temp_docx = os.path.join(temp_dir, 'temp.docx')
                subprocess.run(['soffice', '--headless', '--convert-to', 'docx', '--outdir', temp_dir, input_path])
                doc = docx.Document(temp_docx)
                text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
            # Translate the content
            try:
                translator = GoogleTranslator(source=source_lang, target=target_lang)
                translated_text = translator.translate(text_content)
            except Exception as e:
                current_app.logger.error(f"Translation error: {str(e)}")
                return jsonify({'error': f"Translation failed: {str(e)}"}), 500
            
            # Create output file in original format
            output_path = os.path.join(temp_dir, f'translated_{filename}')
            
            try:
                if file_extension == '.pdf':
                    # Create PDF with translated content
                    c = canvas.Canvas(output_path, pagesize=letter)
                    width, height = letter
                    c.setFont('Helvetica', 12)
                    y = height - 50
                    
                    for line in translated_text.split('\n'):
                        if y < 50:
                            c.showPage()
                            c.setFont('Helvetica', 12)
                            y = height - 50
                        c.drawString(50, y, line)
                        y -= 15
                    c.save()
                    
                elif file_extension == '.docx':
                    doc = docx.Document()
                    for para in translated_text.split('\n'):
                        if para.strip():
                            doc.add_paragraph(para)
                    doc.save(output_path)
                    
                else:  # .txt and others
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(translated_text)
                
                return send_file(
                    output_path,
                    as_attachment=True,
                    download_name=f'translated_{filename}'
                )
                
            except Exception as e:
                current_app.logger.error(f"Error creating output file: {str(e)}")
                return jsonify({'error': f"Error creating output file: {str(e)}"}), 500
                
    except Exception as e:
        current_app.logger.error(f"Translation error: {str(e)}")
        return jsonify({'error': f"Translation failed: {str(e)}"}), 500
